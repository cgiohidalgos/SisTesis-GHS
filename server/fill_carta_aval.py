import sys, json
from docx import Document
from docx.shared import Inches

data = json.loads(sys.argv[1])
template = sys.argv[2]
output = sys.argv[3]

directors_list = [d.strip() for d in data.get('director', '').split(' y ') if d.strip()]
multiple_directors = len(directors_list) > 1

programa = data.get('programa', '')
replacements = {
    'Nombre trabajo de grado': data.get('titulo', ''),
    'NombrePrograma': programa,
    'de Programa de Ingeniería de Sistemas': f'de Programa de {programa}' if programa else 'de Programa',
    'NombreEstudiante': data.get('estudiante', ''),
    'CodigoEstudiante': data.get('codigo', ''),
    'NombreDirectorProyecto': data.get('director', ''),
    'Agradecemos la atención a la presente.': 'Agradecemos la atención a la presente.' if multiple_directors else 'Agradezco la atención a la presente.',
}

doc = Document(template)

def replace_in_para(para):
    for run in para.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

for para in doc.paragraphs:
    replace_in_para(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_para(para)

from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

for section in doc.sections:
    section.top_margin    = Inches(1.8)
    section.bottom_margin = Inches(0.8)

# Reduce paragraph spacing to fit on one page
for para in doc.paragraphs:
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

doc.save(output)
