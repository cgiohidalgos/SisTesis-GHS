#!/usr/bin/env python3
"""
Genera los 4 manuales de usuario de SisTesis en formato .docx
con encabezado y pie de página institucional USB Cali.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
HEADER_IMG = os.path.join(BASE_DIR, "Formatos", "header.png")
FOOTER_IMG = os.path.join(BASE_DIR, "Formatos", "footer.png")
OUTPUT_DIR = os.path.join(BASE_DIR, "Manuales")
os.makedirs(OUTPUT_DIR, exist_ok=True)

URL = "https://sistesis.site"


# ─── Helpers ───────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header image
    if os.path.exists(HEADER_IMG):
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run()
        run.add_picture(HEADER_IMG, width=Inches(6.5))

    # Footer image
    if os.path.exists(FOOTER_IMG):
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.add_picture(FOOTER_IMG, width=Inches(6.5))

    return doc


def title_page(doc, title, subtitle=""):
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTESIS")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC0, 0x5A, 0x12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Sistema de Gestión de Proyectos de Grado")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(title)
    run3.bold = True
    run3.font.size = Pt(22)
    run3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if subtitle:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(subtitle)
        run4.font.size = Pt(13)
        run4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    for _ in range(3):
        doc.add_paragraph("")

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Facultad de Ingeniería")
    run5.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("Universidad de San Buenaventura – Cali")
    run6.font.size = Pt(12)

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run7 = p7.add_run(URL)
    run7.font.size = Pt(11)
    run7.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    doc.add_page_break()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x5A, 0x12) if level == 1 else RGBColor(0x33, 0x33, 0x33)
    return h


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p


def numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p


def table_simple(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph("")
    return table


def note_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("ℹ️  " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)
    return p


def warning_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠️  " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    return p


def save(doc, name):
    path = os.path.join(OUTPUT_DIR, name)
    doc.save(path)
    print(f"  ✓ {path}")


# ═══════════════════════════════════════════════════════════════════
#  MANUAL 1: ESTUDIANTE
# ═══════════════════════════════════════════════════════════════════

def manual_estudiante():
    doc = new_doc()
    title_page(doc, "Manual de Usuario", "Rol: Estudiante")

    # ── Índice ──
    heading(doc, "Tabla de Contenido")
    toc_items = [
        "1. Introducción",
        "2. Requisitos previos",
        "3. Registro en el sistema",
        "4. Inicio de sesión",
        "5. Panel principal (Dashboard)",
        "6. Registrar un proyecto de grado",
        "7. Enviar proyecto a evaluación",
        "8. Seguimiento del proyecto (Timeline)",
        "9. Enviar correcciones",
        "10. Firma de documentos",
        "11. Preguntas frecuentes",
    ]
    for item in toc_items:
        bullet(doc, item)
    doc.add_page_break()

    # ── 1. Introducción ──
    heading(doc, "1. Introducción")
    para(doc, "SisTesis es el Sistema de Gestión de Proyectos de Grado de la Facultad de Ingeniería de la Universidad de San Buenaventura – Cali. Este manual describe el uso de la plataforma desde el rol de Estudiante.")
    para(doc, "Como estudiante, usted podrá:")
    bullet(doc, "Registrar su proyecto de grado con toda la información requerida.")
    bullet(doc, "Subir el documento de tesis y archivos complementarios.")
    bullet(doc, "Enviar el proyecto para evaluación académica.")
    bullet(doc, "Hacer seguimiento del progreso de su proyecto en tiempo real.")
    bullet(doc, "Enviar correcciones cuando sean solicitadas por los evaluadores.")
    bullet(doc, "Firmar el acta de sustentación de forma digital.")
    doc.add_page_break()

    # ── 2. Requisitos previos ──
    heading(doc, "2. Requisitos previos")
    bullet(doc, "Navegador web actualizado (Google Chrome, Firefox, Edge o Safari).")
    bullet(doc, "Correo electrónico institucional activo (@usbcali.edu.co).")
    bullet(doc, "Documento de tesis en formato PDF o DOCX (máximo 100 MB).")
    bullet(doc, "Conexión a internet estable.")
    para(doc, "")
    note_box(doc, f"Acceda al sistema en: {URL}")
    doc.add_page_break()

    # ── 3. Registro ──
    heading(doc, "3. Registro en el sistema")
    para(doc, "Si es la primera vez que utiliza SisTesis, debe crear una cuenta de estudiante.")
    heading(doc, "Pasos para registrarse", level=2)
    numbered(doc, f"Ingrese a {URL}/register/student en su navegador.")
    numbered(doc, "Complete el formulario con los siguientes datos:")
    bullet(doc, "Nombre completo", level=1)
    bullet(doc, "Código de estudiante", level=1)
    bullet(doc, "Cédula de ciudadanía", level=1)
    bullet(doc, "Correo electrónico institucional", level=1)
    bullet(doc, "Contraseña (mínimo 6 caracteres)", level=1)
    numbered(doc, 'Haga clic en el botón "Registrarse".')
    numbered(doc, "Recibirá un correo de confirmación con sus credenciales.")
    para(doc, "")
    warning_box(doc, "Use su correo institucional. El sistema puede rechazar correos personales.")
    doc.add_page_break()

    # ── 4. Inicio de sesión ──
    heading(doc, "4. Inicio de sesión")
    numbered(doc, f"Ingrese a {URL}/login/student.")
    numbered(doc, "Escriba su correo institucional (o código/cédula) y contraseña.")
    numbered(doc, 'Haga clic en "Iniciar sesión".')
    numbered(doc, "Será redirigido al panel principal del estudiante.")
    para(doc, "")
    note_box(doc, "Si olvidó su contraseña, contacte al administrador del programa para que le envíe nuevas credenciales.")
    doc.add_page_break()

    # ── 5. Panel principal ──
    heading(doc, "5. Panel principal (Dashboard)")
    para(doc, "Al iniciar sesión, verá el panel principal del estudiante. Desde aquí puede:")
    bullet(doc, "Ver el estado actual de su proyecto de grado.")
    bullet(doc, "Acceder a las acciones disponibles según el estado del proyecto.")
    bullet(doc, "Navegar al timeline (línea de tiempo) del proyecto.")

    heading(doc, "Información mostrada", level=2)
    table_simple(doc,
        ["Campo", "Descripción"],
        [
            ["Título del proyecto", "Nombre completo del proyecto de grado"],
            ["Estudiante(s)", "Nombre(s) del o los autores"],
            ["Estado", "Estado actual del proyecto (borrador, enviado, en evaluación, etc.)"],
            ["Fecha de envío", "Fecha en que se envió el proyecto a evaluación"],
            ["Evaluadores", "Nombre de los evaluadores asignados (si aplica)"],
        ]
    )

    heading(doc, "Acciones disponibles (estado Borrador)", level=2)
    bullet(doc, "Enviar a evaluar: Envía el proyecto al administrador para revisión.")
    bullet(doc, "Modificar: Permite editar los datos del proyecto.")
    bullet(doc, "Eliminar proyecto de grado: Elimina permanentemente el proyecto.")
    para(doc, "")
    warning_box(doc, "Una vez enviado a evaluar, no podrá modificar ni eliminar el proyecto.")
    doc.add_page_break()

    # ── 6. Registrar proyecto ──
    heading(doc, "6. Registrar un proyecto de grado")
    para(doc, "Para crear un nuevo proyecto de grado, siga estos pasos:")

    heading(doc, "Paso a paso", level=2)
    numbered(doc, 'En el panel principal, haga clic en "Registrar nuevo proyecto de grado".')
    numbered(doc, "Complete el formulario con la siguiente información:")
    para(doc, "")

    table_simple(doc,
        ["Campo", "Obligatorio", "Descripción"],
        [
            ["Título del proyecto", "Sí", "Nombre completo del proyecto de grado"],
            ["Resumen (Abstract)", "Sí", "Descripción resumida del proyecto"],
            ["Palabras clave", "No", "Términos clave de investigación separados por coma"],
            ["Documento de tesis", "Sí", "Archivo PDF o DOCX del documento (máx. 100 MB)"],
            ["Archivo de aval", "No", "Documento de aval del director"],
            ["URL externa", "No", "Enlace a repositorio GitHub u otro recurso"],
            ["Programa(s)", "Sí", "Seleccione uno o más programas académicos"],
            ["Director(es)", "No", "Seleccione director(es) propuestos de la lista"],
        ]
    )

    heading(doc, "Agregar coautor (segundo estudiante)", level=2)
    para(doc, "Si el proyecto es en pareja, complete los datos del compañero:")
    bullet(doc, "Nombre completo del compañero")
    bullet(doc, "Código de estudiante")
    bullet(doc, "Cédula de ciudadanía")
    bullet(doc, "Correo institucional")

    numbered(doc, 'Haga clic en "Registrar" para guardar el proyecto como borrador.')
    numbered(doc, "El proyecto aparecerá en su panel principal con estado Borrador.")
    para(doc, "")
    note_box(doc, "Formatos de archivo permitidos: PDF, DOC, DOCX, JPG, PNG, GIF. Máximo 5 archivos por carga.")
    warning_box(doc, "Verifique que las ventanas de recepción del programa estén abiertas. No podrá enviar a un programa con recepción cerrada.")
    doc.add_page_break()

    # ── 7. Enviar a evaluación ──
    heading(doc, "7. Enviar proyecto a evaluación")
    para(doc, "Cuando su proyecto esté completo y listo para ser evaluado:")
    numbered(doc, "Vaya al panel principal del estudiante.")
    numbered(doc, "Verifique que toda la información y archivos estén correctos.")
    numbered(doc, 'Haga clic en el botón "Enviar a evaluar".')
    numbered(doc, "Confirme el envío en el diálogo de confirmación.")
    numbered(doc, "El estado cambiará a Enviado.")
    para(doc, "")
    para(doc, "A partir de este momento:")
    bullet(doc, "El administrador del programa revisará su documentación.")
    bullet(doc, "Se le asignarán evaluadores.")
    bullet(doc, "Recibirá notificaciones por correo sobre el avance.")
    warning_box(doc, "Esta acción es irreversible. Asegúrese de que su documento esté completo antes de enviarlo.")
    doc.add_page_break()

    # ── 8. Seguimiento ──
    heading(doc, "8. Seguimiento del proyecto (Timeline)")
    para(doc, "La línea de tiempo le permite ver el progreso completo de su proyecto de grado.")

    heading(doc, "Cómo acceder", level=2)
    numbered(doc, "En el panel principal, haga clic en su proyecto.")
    numbered(doc, "Navegue a la sección Timeline / Línea de tiempo.")

    heading(doc, "Eventos que verá en la línea de tiempo", level=2)
    table_simple(doc,
        ["Evento", "Descripción"],
        [
            ["Proyecto enviado", "Fecha en que envió su proyecto"],
            ["Revisión administrativa", "El administrador revisó su documentación"],
            ["Evaluadores asignados", "Se asignaron evaluadores a su proyecto"],
            ["En evaluación académica", "Los evaluadores están revisando su trabajo"],
            ["Concepto emitido", "Los evaluadores emitieron su concepto"],
            ["Correcciones solicitadas", "Se requieren cambios en su documento"],
            ["Aprobada para sustentación", "Su proyecto fue aprobado para sustentar"],
            ["Sustentación programada", "Fecha, hora y lugar de la sustentación"],
            ["Sustentación evaluada", "Los evaluadores calificaron su sustentación"],
            ["Acta firmada", "Todas las firmas fueron recolectadas"],
            ["Proceso finalizado", "Su proyecto de grado fue completado"],
        ]
    )

    heading(doc, "Información adicional en el timeline", level=2)
    bullet(doc, "Calificaciones parciales y finales (cuando estén disponibles).")
    bullet(doc, "Observaciones y retroalimentación de los evaluadores.")
    bullet(doc, "Archivos adjuntos de retroalimentación.")
    bullet(doc, "Fecha y lugar de la sustentación.")
    doc.add_page_break()

    # ── 9. Correcciones ──
    heading(doc, "9. Enviar correcciones")
    para(doc, "Si los evaluadores solicitan correcciones a su documento:")

    numbered(doc, "Recibirá una notificación por correo electrónico.")
    numbered(doc, "Ingrese al sistema y vaya a la línea de tiempo de su proyecto.")
    numbered(doc, "Revise las observaciones y retroalimentación de cada evaluador.")
    numbered(doc, "Prepare la versión corregida de su documento.")
    numbered(doc, "En la sección de correcciones del timeline:")
    bullet(doc, "Escriba un comentario explicando los cambios realizados.", level=1)
    bullet(doc, "Adjunte el documento corregido (PDF/DOCX).", level=1)
    numbered(doc, 'Haga clic en "Enviar correcciones".')
    numbered(doc, "El sistema incrementará la ronda de revisión.")
    numbered(doc, "Los evaluadores recibirán notificación para re-evaluar.")
    para(doc, "")
    note_box(doc, "Puede haber múltiples rondas de corrección. Cada ronda queda registrada en la línea de tiempo.")
    doc.add_page_break()

    # ── 10. Firma ──
    heading(doc, "10. Firma de documentos")
    para(doc, "Al finalizar el proceso de evaluación, se generará un acta de sustentación que debe ser firmada digitalmente.")

    heading(doc, "Proceso de firma", level=2)
    numbered(doc, "Recibirá un correo electrónico con un enlace único de firma.")
    numbered(doc, "Haga clic en el enlace (no requiere iniciar sesión).")
    numbered(doc, "En la página de firma:")
    bullet(doc, "Revise la información del proyecto y el acta.", level=1)
    bullet(doc, 'Haga clic en "Descargar PDF" para revisar el documento completo.', level=1)
    numbered(doc, "Elija un método de firma:")
    bullet(doc, "Dibujar: Use el mouse o pantalla táctil para dibujar su firma.", level=1)
    bullet(doc, "Subir imagen: Cargue una imagen escaneada de su firma.", level=1)
    numbered(doc, "Suba el PDF firmado si es requerido.")
    numbered(doc, 'Haga clic en "Enviar firma".')
    numbered(doc, "Será redirigido a una página de confirmación.")
    para(doc, "")
    note_box(doc, "El enlace de firma es único y personal. No lo comparta con otras personas.")

    heading(doc, "Carta meritoria", level=2)
    para(doc, "Si su proyecto obtuvo una calificación meritoria (≥ 4.8/5.0), recibirá un enlace adicional para firmar la carta de graduación meritoria. El proceso es idéntico al de firma del acta.")
    doc.add_page_break()

    # ── 11. FAQ ──
    heading(doc, "11. Preguntas frecuentes")

    heading(doc, "¿Puedo editar mi proyecto después de enviarlo?", level=2)
    para(doc, "No. Una vez enviado a evaluación, el proyecto queda en modo solo lectura. Asegúrese de que toda la información esté correcta antes de enviarlo.")

    heading(doc, "¿Cuántos archivos puedo subir?", level=2)
    para(doc, "Puede subir hasta 5 archivos por carga, con un tamaño máximo de 100 MB cada uno. Formatos: PDF, DOC, DOCX, JPG, PNG, GIF.")

    heading(doc, "¿Cómo sé si mis evaluadores ya revisaron mi proyecto?", level=2)
    para(doc, "Revise la línea de tiempo de su proyecto. Cuando un evaluador emita su concepto, aparecerá un evento nuevo con la retroalimentación.")

    heading(doc, "¿Qué significan las calificaciones?", level=2)
    table_simple(doc,
        ["Calificación", "Resultado"],
        [
            ["≥ 4.8 / 5.0", "Aprobada Meritoria"],
            ["≥ 3.0 / 5.0", "Aprobada"],
            ["< 3.0 / 5.0", "No Aprobada"],
        ]
    )
    para(doc, "La calificación final se calcula como: (Nota Documento × Peso%) + (Nota Sustentación × Peso%). Los pesos por defecto son 70% documento y 30% sustentación.")

    heading(doc, "¿Qué hago si no puedo acceder?", level=2)
    para(doc, "Contacte al administrador de su programa académico para restablecer su contraseña o verificar su cuenta.")

    save(doc, "Manual_Estudiante.docx")


# ═══════════════════════════════════════════════════════════════════
#  MANUAL 2: EVALUADOR
# ═══════════════════════════════════════════════════════════════════

def manual_evaluador():
    doc = new_doc()
    title_page(doc, "Manual de Usuario", "Rol: Evaluador")

    heading(doc, "Tabla de Contenido")
    toc = [
        "1. Introducción",
        "2. Inicio de sesión",
        "3. Panel del evaluador",
        "4. Evaluación de documento",
        "5. Rúbrica de evaluación del documento",
        "6. Evaluación de sustentación",
        "7. Rúbrica de evaluación de sustentación",
        "8. Concepto y observaciones",
        "9. Evaluación de correcciones (rondas múltiples)",
        "10. Firma del acta",
        "11. Preguntas frecuentes",
    ]
    for item in toc:
        bullet(doc, item)
    doc.add_page_break()

    # ── 1 ──
    heading(doc, "1. Introducción")
    para(doc, "SisTesis es el Sistema de Gestión de Proyectos de Grado de la Facultad de Ingeniería de la Universidad de San Buenaventura – Cali. Este manual describe el uso de la plataforma desde el rol de Evaluador.")
    para(doc, "Como evaluador, usted podrá:")
    bullet(doc, "Ver los proyectos de grado que le han sido asignados para evaluación.")
    bullet(doc, "Evaluar el documento de tesis mediante una rúbrica detallada.")
    bullet(doc, "Evaluar la sustentación oral del proyecto (cuando aplique).")
    bullet(doc, "Emitir un concepto académico (Aceptado, Cambios menores, Cambios mayores).")
    bullet(doc, "Re-evaluar en rondas de corrección si el estudiante envía revisiones.")
    bullet(doc, "Firmar digitalmente el acta de sustentación.")
    doc.add_page_break()

    # ── 2 ──
    heading(doc, "2. Inicio de sesión")
    numbered(doc, f"Ingrese a {URL}/login/staff en su navegador.")
    numbered(doc, "Escriba su correo institucional y contraseña.")
    numbered(doc, 'Haga clic en "Iniciar sesión".')
    numbered(doc, "Será redirigido al panel del evaluador.")
    para(doc, "")
    note_box(doc, "Sus credenciales le fueron enviadas por correo electrónico al ser registrado como evaluador. Si no las tiene, contacte al administrador del programa.")
    doc.add_page_break()

    # ── 3 ──
    heading(doc, "3. Panel del evaluador")
    para(doc, "Al ingresar, verá una lista de los proyectos de grado asignados a usted.")
    heading(doc, "Información de cada proyecto", level=2)
    table_simple(doc,
        ["Elemento", "Descripción"],
        [
            ["Título", "Nombre del proyecto de grado"],
            ["Estudiante(s)", "Nombre(s) del o los autores"],
            ["Estado", "Estado actual del proyecto"],
            ["Ronda de revisión", "Número de la ronda actual (si hay correcciones)"],
            ["Indicador de evaluación", "Muestra si ya evaluó esta ronda"],
            ["Indicador de acta", "Muestra si el acta está lista para firmar"],
        ]
    )
    para(doc, "Haga clic en cualquier proyecto para acceder a la interfaz de evaluación.")
    doc.add_page_break()

    # ── 4 ──
    heading(doc, "4. Evaluación de documento")
    para(doc, "Al hacer clic en un proyecto, accederá a la interfaz completa de evaluación.")

    heading(doc, "Sección superior: Información del proyecto", level=2)
    bullet(doc, "Título del proyecto (con indicador de ronda de revisión si aplica).")
    bullet(doc, "Datos de los estudiantes: nombre, código, cédula, correo.")
    bullet(doc, "Archivos adjuntos: enlaces para descargar el documento de tesis y archivos complementarios.")
    bullet(doc, "Información de sustentación: fecha, lugar y datos adicionales (si está programada).")
    bullet(doc, "Fecha límite de evaluación asignada por el administrador.")

    heading(doc, "Pasos para evaluar el documento", level=2)
    numbered(doc, "Descargue y lea el documento de tesis completo.")
    numbered(doc, "Navegue a la pestaña/sección de Evaluación del Documento.")
    numbered(doc, "Complete la rúbrica asignando una calificación de 0 a 5 para cada criterio.")
    numbered(doc, "Escriba observaciones específicas para cada criterio evaluado.")
    numbered(doc, "Redacte las observaciones generales del documento.")
    numbered(doc, "Seleccione el concepto general (ver sección 8).")
    numbered(doc, "Opcionalmente, adjunte archivos de retroalimentación.")
    numbered(doc, 'Haga clic en "Enviar Evaluación".')
    doc.add_page_break()

    # ── 5 ──
    heading(doc, "5. Rúbrica de evaluación del documento")
    para(doc, "La rúbrica del documento está organizada en secciones con pesos porcentuales. Cada criterio se califica de 0 a 5.")

    table_simple(doc,
        ["Sección", "Peso", "Criterios"],
        [
            ["Fundamentación Problémica", "30%", "Definición del problema, Justificación, Marco teórico, Estado del arte, Referencias IEEE"],
            ["Propuesta Metodológica", "30%", "Objetivo general, Objetivos específicos, Coherencia metodológica, Cumplimiento"],
            ["Aspectos Disciplinares", "30%", "Implementación técnica, Métricas y validación, Análisis de resultados, Discusión técnica"],
            ["Presentación del Documento", "10%", "Redacción, Cumplimiento de normas"],
        ]
    )

    para(doc, "La calificación del documento se calcula automáticamente como el promedio ponderado de todas las secciones, en una escala de 0 a 5.")
    note_box(doc, "La rúbrica puede variar según el programa académico. El administrador puede personalizar las secciones, pesos y criterios.")
    doc.add_page_break()

    # ── 6 ──
    heading(doc, "6. Evaluación de sustentación")
    para(doc, "La evaluación de sustentación solo aparece cuando el administrador ha programado una fecha de defensa.")

    heading(doc, "Pasos para evaluar la sustentación", level=2)
    numbered(doc, "Asista a la sustentación en la fecha, hora y lugar indicados.")
    numbered(doc, "Después de la sustentación, ingrese al sistema.")
    numbered(doc, "Navegue a la pestaña/sección de Evaluación de Sustentación.")
    numbered(doc, "Complete la rúbrica de sustentación con calificaciones de 0 a 5.")
    numbered(doc, "Escriba observaciones sobre la presentación oral.")
    numbered(doc, "Seleccione el concepto.")
    numbered(doc, 'Haga clic en "Enviar Evaluación".')
    doc.add_page_break()

    # ── 7 ──
    heading(doc, "7. Rúbrica de evaluación de sustentación")
    table_simple(doc,
        ["Sección", "Peso", "Criterios"],
        [
            ["Claridad y Dominio del Problema", "25%", "Dominio del tema, claridad en la exposición del problema"],
            ["Dominio Metodológico", "25%", "Comprensión y explicación de la metodología empleada"],
            ["Dominio Técnico y Resultados", "30%", "Presentación de implementación, resultados y métricas"],
            ["Comunicación y Presentación", "20%", "Calidad de las diapositivas, manejo del tiempo, respuesta a preguntas"],
        ]
    )
    doc.add_page_break()

    # ── 8 ──
    heading(doc, "8. Concepto y observaciones")
    para(doc, "Al finalizar la evaluación (documento o sustentación), debe seleccionar uno de los siguientes conceptos:")

    table_simple(doc,
        ["Concepto", "Significado"],
        [
            ["Aceptado", "El trabajo cumple con los requisitos. No requiere cambios."],
            ["Cambios menores", "El trabajo requiere ajustes pequeños que no afectan la estructura."],
            ["Cambios mayores", "El trabajo requiere modificaciones significativas y debe ser re-evaluado."],
        ]
    )

    heading(doc, "Cálculo de calificación final", level=2)
    para(doc, "El sistema calcula automáticamente:")
    bullet(doc, "Nota del documento: Promedio ponderado de la rúbrica del documento.")
    bullet(doc, "Nota de sustentación: Promedio ponderado de la rúbrica de sustentación.")
    bullet(doc, "Nota final ponderada = (Documento × Peso%) + (Sustentación × Peso%).")
    para(doc, "Los pesos por defecto son 70% documento y 30% sustentación, pero pueden variar por programa.")

    heading(doc, "Clasificación del resultado", level=2)
    table_simple(doc,
        ["Rango", "Clasificación"],
        [
            ["≥ 4.8 / 5.0", "APROBADA MERITORIA"],
            ["≥ 3.0 / 5.0", "APROBADA"],
            ["< 3.0 / 5.0", "NO APROBADA"],
        ]
    )
    doc.add_page_break()

    # ── 9 ──
    heading(doc, "9. Evaluación de correcciones (rondas múltiples)")
    para(doc, "Si el concepto emitido requiere cambios, el estudiante enviará una versión corregida.")

    numbered(doc, "Recibirá una notificación cuando el estudiante envíe correcciones.")
    numbered(doc, "El sistema incrementará la ronda de revisión (Ronda 1, 2, 3...).")
    numbered(doc, "Ingrese al proyecto y verá la nueva ronda con los cambios del estudiante.")
    numbered(doc, "Descargue el documento revisado y los comentarios del estudiante.")
    numbered(doc, "Evalúe nuevamente completando la rúbrica de la nueva ronda.")
    numbered(doc, "Las evaluaciones de rondas anteriores quedan bloqueadas pero visibles.")
    para(doc, "")
    note_box(doc, "Puede haber múltiples rondas de corrección hasta que el trabajo sea aceptado.")
    doc.add_page_break()

    # ── 10 ──
    heading(doc, "10. Firma del acta")
    para(doc, "Una vez completadas todas las evaluaciones y aprobado el proyecto:")

    numbered(doc, "El sistema mostrará el estado 'Acta preparada' en la sección de firmas.")
    numbered(doc, "Puede firmar directamente desde la interfaz de evaluación, o a través de un enlace enviado por correo.")
    numbered(doc, "Para firmar:")
    bullet(doc, 'Opción 1: Haga clic en "Dibujar firma" y use el mouse/pantalla táctil.', level=1)
    bullet(doc, 'Opción 2: Haga clic en "Cargar firma" y suba una imagen de su firma.', level=1)
    numbered(doc, 'Haga clic en "Guardar firma".')
    numbered(doc, "El estado de firma se actualizará como firmado con la fecha.")
    para(doc, "")
    note_box(doc, "La sección de firma muestra el estado de todos los firmantes: evaluadores, director y director de programa.")
    doc.add_page_break()

    # ── 11 ──
    heading(doc, "11. Preguntas frecuentes")

    heading(doc, "¿Puedo modificar una evaluación ya enviada?", level=2)
    para(doc, "Solo puede modificar evaluaciones de la ronda activa. Las evaluaciones de rondas anteriores quedan bloqueadas.")

    heading(doc, "¿Qué pasa si no alcanzo a evaluar antes de la fecha límite?", level=2)
    para(doc, "La evaluación aparecerá como vencida en el panel del administrador. Contacte al administrador si necesita una extensión.")

    heading(doc, "¿Puedo ver las evaluaciones de otros evaluadores?", level=2)
    para(doc, "No directamente. La evaluación es independiente. Sin embargo, en la línea de tiempo puede ver eventos generales del proyecto.")

    heading(doc, "¿La asignación puede ser a ciegas?", level=2)
    para(doc, "Sí. Si el administrador activa la asignación ciega (pares ciegos), usted no verá los nombres de los otros evaluadores y los estudiantes no verán su nombre.")

    save(doc, "Manual_Evaluador.docx")


# ═══════════════════════════════════════════════════════════════════
#  MANUAL 3: PROFESOR ADMINISTRADOR
# ═══════════════════════════════════════════════════════════════════

def manual_admin():
    doc = new_doc()
    title_page(doc, "Manual de Usuario", "Rol: Profesor Administrador")

    heading(doc, "Tabla de Contenido")
    toc = [
        "1. Introducción",
        "2. Inicio de sesión",
        "3. Panel de administración (Dashboard)",
        "4. Gestión de proyectos de grado",
        "5. Detalle de un proyecto",
        "6. Asignación de evaluadores",
        "7. Seguimiento de evaluaciones",
        "8. Gestión de correcciones y rondas de revisión",
        "9. Programar sustentación",
        "10. Acta de sustentación y firmas digitales",
        "11. Finalización del proyecto",
        "12. Gestión de evaluadores",
        "13. Gestión de programas académicos",
        "14. Configuración de rúbricas",
        "15. Elementos de revisión (checklist)",
        "16. Pesos de evaluación",
        "17. Historial de notificaciones",
        "18. Exportar datos (CSV)",
        "19. Preguntas frecuentes",
    ]
    for item in toc:
        bullet(doc, item)
    doc.add_page_break()

    # ── 1 ──
    heading(doc, "1. Introducción")
    para(doc, "SisTesis es el Sistema de Gestión de Proyectos de Grado de la Facultad de Ingeniería de la Universidad de San Buenaventura – Cali. Este manual describe el uso de la plataforma desde el rol de Profesor Administrador.")
    para(doc, "Como administrador de programa, usted podrá:")
    bullet(doc, "Revisar y gestionar los proyectos de grado de sus programas asignados.")
    bullet(doc, "Asignar evaluadores y establecer fechas límite.")
    bullet(doc, "Monitorear el progreso de evaluaciones en tiempo real.")
    bullet(doc, "Solicitar correcciones y gestionar rondas de revisión.")
    bullet(doc, "Programar sustentaciones y gestionar el acta de firmas.")
    bullet(doc, "Configurar rúbricas, pesos de evaluación y elementos de revisión.")
    bullet(doc, "Gestionar la base de datos de evaluadores.")
    bullet(doc, "Administrar programas académicos y ventanas de recepción.")
    bullet(doc, "Exportar informes en formato CSV.")
    doc.add_page_break()

    # ── 2 ──
    heading(doc, "2. Inicio de sesión")
    numbered(doc, f"Ingrese a {URL}/login/staff en su navegador.")
    numbered(doc, "Escriba su correo institucional y contraseña.")
    numbered(doc, 'Haga clic en "Iniciar sesión".')
    numbered(doc, "Será redirigido al panel de administración.")
    note_box(doc, "Las credenciales de administrador son otorgadas por el Superadministrador del sistema.")
    doc.add_page_break()

    # ── 3 ──
    heading(doc, "3. Panel de administración (Dashboard)")
    para(doc, "El panel principal muestra un resumen general del estado de los proyectos de sus programas.")

    heading(doc, "Indicadores estadísticos", level=2)
    table_simple(doc,
        ["Indicador", "Descripción"],
        [
            ["Total proyectos", "Cantidad total de proyectos registrados"],
            ["En evaluación", "Proyectos actualmente en proceso de evaluación"],
            ["Finalizadas", "Proyectos que completaron todo el proceso"],
            ["Evaluadores", "Cantidad de evaluadores registrados"],
            ["Evaluaciones vencidas", "Evaluaciones que superaron la fecha límite (en rojo)"],
            ["Vence < 7 días", "Evaluaciones próximas a vencer (en amarillo)"],
            ["Vence < 15 días", "Evaluaciones con vencimiento medio"],
            ["Vence < 30 días", "Evaluaciones con plazo amplio"],
        ]
    )

    heading(doc, "Gráficos", level=2)
    bullet(doc, "Gráfico circular: Distribución de proyectos por estado.")
    bullet(doc, "Gráfico de barras: Proyectos por programa (en evaluación, finalizados, otros).")
    bullet(doc, "Tabla de evaluadores: Asignaciones, completadas, pendientes y vencimientos por evaluador.")

    heading(doc, "Menú lateral de navegación", level=2)
    table_simple(doc,
        ["Opción", "Descripción"],
        [
            ["Panel", "Dashboard con estadísticas"],
            ["Proyectos", "Lista de todos los proyectos de grado"],
            ["Evaluadores", "Gestión de cuentas de evaluadores"],
            ["Evaluaciones", "Monitor de evaluaciones y vencimientos"],
            ["Programas", "Gestión de programas académicos"],
            ["Rúbricas", "Configuración de rúbricas de evaluación"],
            ["Elementos de Revisión", "Checklist administrativo por programa"],
            ["Pesos de Evaluación", "Configurar porcentajes documento/sustentación"],
            ["Notificaciones", "Historial de correos enviados"],
        ]
    )
    doc.add_page_break()

    # ── 4 ──
    heading(doc, "4. Gestión de proyectos de grado")
    para(doc, 'Desde el menú lateral, seleccione "Proyectos" para ver la lista completa.')

    heading(doc, "Vista de lista", level=2)
    bullet(doc, "Cada proyecto muestra: título, estudiantes, estado actual y evaluadores asignados.")
    bullet(doc, "Badge de estado con código de color para identificar rápidamente.")
    bullet(doc, "Botón de eliminar en cada tarjeta (con confirmación).")
    bullet(doc, 'Botón "Exportar CSV" para descargar la lista completa.')

    heading(doc, "Acciones", level=2)
    bullet(doc, "Haga clic en un proyecto para ver su detalle completo.")
    bullet(doc, "Use los filtros de estado para encontrar proyectos específicos.")
    doc.add_page_break()

    # ── 5 ──
    heading(doc, "5. Detalle de un proyecto")
    para(doc, "La vista de detalle es la interfaz principal para gestionar un proyecto individual.")

    heading(doc, "Secciones disponibles", level=2)

    heading(doc, "A. Información general", level=3)
    bullet(doc, "Título completo del proyecto (con ronda de revisión si aplica).")
    bullet(doc, "Datos de estudiantes: nombre, código, cédula, correo institucional.")
    bullet(doc, "Evaluadores asignados con fecha límite.")
    bullet(doc, "Directores propuestos.")
    bullet(doc, "Archivos adjuntos (descargables).")
    bullet(doc, "Estado actual con badge de color.")

    heading(doc, "B. Acciones administrativas", level=3)
    bullet(doc, "Cambiar estado: Menú desplegable para cambiar manualmente el estado del proyecto.")
    bullet(doc, "Asignar evaluadores: Seleccionar evaluadores, fecha límite y modo ciego.")
    bullet(doc, "Reemplazar evaluador: Sustituir un evaluador por otro.")
    bullet(doc, "Eliminar evaluador: Remover un evaluador de la asignación.")
    bullet(doc, "Programar sustentación: Fecha, hora, lugar e información adicional.")

    heading(doc, "C. Revisión documental", level=3)
    bullet(doc, "Checklist de elementos de revisión (configurables por programa).")
    bullet(doc, "Campo de retroalimentación para el estudiante.")
    bullet(doc, "Botón para adjuntar archivos de observaciones.")

    heading(doc, "D. Calificaciones y resultados", level=3)
    bullet(doc, "Calificaciones individuales de cada evaluador (por ronda).")
    bullet(doc, "Promedios consolidados: documento, sustentación, nota final ponderada.")
    bullet(doc, "Clasificación: APROBADA, APROBADA MERITORIA o NO APROBADA.")
    bullet(doc, "Opción para sobrescribir la nota final ponderada manualmente.")

    heading(doc, "E. Acta y firmas (ver sección 10)", level=3)
    heading(doc, "F. Carta meritoria (si aplica)", level=3)
    para(doc, "Para proyectos con calificación ≥ 4.8, se genera una carta meritoria con un proceso de firma independiente.")

    heading(doc, "G. Línea de tiempo", level=3)
    para(doc, "Historial completo de eventos del proyecto con fechas, actores y observaciones.")
    doc.add_page_break()

    # ── 6 ──
    heading(doc, "6. Asignación de evaluadores")
    numbered(doc, "Abra el detalle del proyecto.")
    numbered(doc, 'Vaya a la sección "Asignar evaluadores".')
    numbered(doc, "Seleccione 2 evaluadores del menú desplegable.")
    numbered(doc, "Establezca la fecha límite de evaluación.")
    numbered(doc, 'Active "Asignación ciega" si desea que los nombres de los evaluadores no sean visibles para los estudiantes.')
    numbered(doc, 'Haga clic en "Asignar".')
    numbered(doc, "El sistema enviará un correo automático a cada evaluador con los datos del proyecto.")
    para(doc, "")
    note_box(doc, "Para reemplazar un evaluador: seleccione el evaluador actual, elija el reemplazo y confirme. Para eliminar un evaluador de la asignación, use el botón de eliminar junto a su nombre.")
    doc.add_page_break()

    # ── 7 ──
    heading(doc, "7. Seguimiento de evaluaciones")
    para(doc, 'Desde el menú lateral, seleccione "Evaluaciones" para acceder al monitor.')

    heading(doc, "Filtros disponibles", level=2)
    table_simple(doc,
        ["Filtro", "Descripción"],
        [
            ["Todas", "Muestra todas las evaluaciones"],
            ["Vencidas", "Evaluaciones que superaron la fecha límite"],
            ["< 7 días", "Evaluaciones que vencen en menos de una semana"],
            ["< 15 días", "Evaluaciones que vencen en menos de dos semanas"],
            ["< 30 días", "Evaluaciones que vencen en menos de un mes"],
        ]
    )

    heading(doc, "Información mostrada", level=2)
    bullet(doc, "Nombre del proyecto.")
    bullet(doc, "Nombre del evaluador.")
    bullet(doc, "Fecha límite.")
    bullet(doc, "Estado de la evaluación (pendiente, completada, vencida).")
    para(doc, "")
    note_box(doc, "Puede hacer clic en los indicadores de vencimiento del Dashboard para navegar directamente al monitor con el filtro pre-seleccionado.")
    doc.add_page_break()

    # ── 8 ──
    heading(doc, "8. Gestión de correcciones y rondas de revisión")
    para(doc, "Cuando los evaluadores emiten un concepto de 'Cambios mayores' o 'Cambios menores':")

    numbered(doc, "Revise las evaluaciones y observaciones de los evaluadores en el detalle del proyecto.")
    numbered(doc, 'Cambie el estado del proyecto a "Correcciones solicitadas".')
    numbered(doc, "Agregue retroalimentación adicional si lo considera necesario.")
    numbered(doc, "El sistema notificará automáticamente al estudiante por correo.")
    numbered(doc, "El estudiante enviará la versión corregida.")
    numbered(doc, "El sistema incrementará la ronda de revisión (Ronda 1, 2, 3...).")
    numbered(doc, "Los evaluadores serán notificados para re-evaluar.")
    numbered(doc, "Repita el proceso hasta que el proyecto sea aceptado.")
    doc.add_page_break()

    # ── 9 ──
    heading(doc, "9. Programar sustentación")
    para(doc, "Una vez que todas las evaluaciones son favorables:")

    numbered(doc, "Abra el detalle del proyecto.")
    numbered(doc, 'Cambie el estado a "Aprobada para sustentación".')
    numbered(doc, 'Vaya a la sección "Programar sustentación".')
    numbered(doc, "Complete los campos:")
    bullet(doc, "Fecha y hora de la sustentación.", level=1)
    bullet(doc, "Lugar (salón, auditorio, enlace virtual).", level=1)
    bullet(doc, "Información adicional (instrucciones, duración, etc.).", level=1)
    numbered(doc, 'Haga clic en "Guardar".')
    numbered(doc, 'El estado cambiará a "Sustentación programada".')
    numbered(doc, "Se enviarán correos automáticos a estudiantes y evaluadores con los datos de la sustentación.")
    doc.add_page_break()

    # ── 10 ──
    heading(doc, "10. Acta de sustentación y firmas digitales")
    para(doc, "Después de la sustentación y evaluación de la presentación:")

    heading(doc, "Gestión del acta", level=2)
    numbered(doc, 'Vaya a la sección "Acta" del detalle del proyecto.')
    numbered(doc, "Verá el estado de firma de cada participante:")

    table_simple(doc,
        ["Firmante", "Descripción"],
        [
            ["Evaluador 1", "Primer evaluador asignado"],
            ["Evaluador 2", "Segundo evaluador asignado"],
            ["Director", "Director del proyecto"],
            ["Director de programa", "Director del programa académico"],
        ]
    )

    heading(doc, "Generar y enviar enlaces de firma", level=2)
    numbered(doc, 'Haga clic en "Generar enlace" para crear un token único de firma.')
    numbered(doc, "Copie el enlace al portapapeles o envíelo directamente por correo.")
    numbered(doc, "Cada firmante recibirá un enlace personal para firmar digitalmente.")
    numbered(doc, "Los firmantes pueden dibujar su firma o subir una imagen.")
    numbered(doc, "El estado de firma se actualiza en tiempo real.")

    heading(doc, "Acciones adicionales", level=2)
    bullet(doc, "Eliminar firma: Revertir una firma si hubo error (con confirmación).")
    bullet(doc, "Descargar acta final: Una vez todas las firmas estén completas.")
    bullet(doc, "Carta meritoria: Para proyectos con nota ≥ 4.8, gestione un proceso de firma separado.")
    doc.add_page_break()

    # ── 11 ──
    heading(doc, "11. Finalización del proyecto")
    numbered(doc, "Verifique que todas las evaluaciones estén completas.")
    numbered(doc, "Verifique que todas las firmas del acta estén recolectadas.")
    numbered(doc, 'Cambie el estado a "Finalizado".')
    numbered(doc, "El sistema clasificará automáticamente el resultado:")
    table_simple(doc,
        ["Calificación final", "Resultado"],
        [
            ["≥ 4.8 / 5.0", "APROBADA MERITORIA"],
            ["≥ 3.0 / 5.0", "APROBADA"],
            ["< 3.0 / 5.0", "NO APROBADA"],
        ]
    )
    doc.add_page_break()

    # ── 12 ──
    heading(doc, "12. Gestión de evaluadores")
    para(doc, 'Desde el menú lateral, seleccione "Evaluadores".')

    heading(doc, "Crear nuevo evaluador", level=2)
    numbered(doc, 'Haga clic en "Crear evaluador".')
    numbered(doc, "Complete el formulario:")
    bullet(doc, "Nombre completo (obligatorio).", level=1)
    bullet(doc, "Correo institucional (obligatorio).", level=1)
    bullet(doc, "Cédula (obligatorio).", level=1)
    bullet(doc, "Especialidad / Área de conocimiento.", level=1)
    bullet(doc, "Contraseña inicial.", level=1)
    numbered(doc, 'Haga clic en "Guardar".')
    numbered(doc, "El sistema creará la cuenta y enviará las credenciales por correo.")

    heading(doc, "Otras acciones", level=2)
    bullet(doc, "Editar evaluador: Modificar datos y contraseña.")
    bullet(doc, "Ver proyectos asignados: Lista de tesis asignadas con estado y fechas.")
    bullet(doc, "Reenviar credenciales: Enviar nuevamente correo con usuario y contraseña.")
    doc.add_page_break()

    # ── 13 ──
    heading(doc, "13. Gestión de programas académicos")
    para(doc, 'Desde el menú lateral, seleccione "Programas".')

    heading(doc, "Crear o editar programa", level=2)
    table_simple(doc,
        ["Campo", "Descripción"],
        [
            ["Nombre", "Nombre del programa académico"],
            ["Fecha inicio recepción", "Fecha desde la cual los estudiantes pueden enviar proyectos"],
            ["Fecha fin recepción", "Fecha hasta la cual se aceptan envíos"],
            ["Máximo evaluadores", "Número máximo de evaluadores por proyecto (por defecto 2)"],
            ["Administradores", "Profesores administradores asignados al programa"],
        ]
    )
    bullet(doc, "Ocultar/mostrar programa: Controla si el programa aparece disponible para los estudiantes.")
    bullet(doc, "Eliminar programa: Remueve el programa (con confirmación).")
    doc.add_page_break()

    # ── 14 ──
    heading(doc, "14. Configuración de rúbricas")
    para(doc, 'Desde el menú lateral, seleccione "Rúbricas".')

    numbered(doc, "Seleccione el programa académico en el menú desplegable.")
    numbered(doc, "Verá dos secciones: Rúbrica del documento y Rúbrica de sustentación.")
    numbered(doc, "Para cada sección puede:")
    bullet(doc, "Editar el nombre de la sección.", level=1)
    bullet(doc, "Modificar el peso porcentual.", level=1)
    bullet(doc, "Agregar, editar o eliminar criterios.", level=1)
    bullet(doc, "Cambiar la puntuación máxima de cada criterio.", level=1)
    numbered(doc, 'Haga clic en "Guardar" para aplicar los cambios.')
    para(doc, "")
    note_box(doc, 'Puede cargar la rúbrica por defecto del sistema usando el botón "Cargar valores por defecto".')
    doc.add_page_break()

    # ── 15 ──
    heading(doc, "15. Elementos de revisión (checklist)")
    para(doc, 'Desde el menú lateral, seleccione "Elementos de Revisión".')

    para(doc, "Estos son los ítems de verificación que aparecen al revisar la documentación de un proyecto:")
    numbered(doc, "Seleccione el programa.")
    numbered(doc, 'Para agregar: Escriba el texto del ítem y haga clic en "Agregar".')
    numbered(doc, "Para editar: Haga clic en el texto del ítem y modifíquelo.")
    numbered(doc, "Para eliminar: Use el botón de eliminar junto a cada ítem.")
    doc.add_page_break()

    # ── 16 ──
    heading(doc, "16. Pesos de evaluación")
    para(doc, 'Desde el menú lateral, seleccione "Pesos de Evaluación".')

    para(doc, "Configure el porcentaje que representa cada componente en la nota final:")
    table_simple(doc,
        ["Componente", "Peso por defecto"],
        [
            ["Documento", "70%"],
            ["Sustentación", "30%"],
        ]
    )
    para(doc, "La suma ideal de ambos pesos debe ser 100%. Los cambios aplican al programa seleccionado.")
    doc.add_page_break()

    # ── 17 ──
    heading(doc, "17. Historial de notificaciones")
    para(doc, 'Desde el menú lateral, seleccione "Notificaciones".')

    heading(doc, "Información mostrada", level=2)
    bullet(doc, "Resumen: total enviadas, fallidas y total general.")
    bullet(doc, "Filtros por estado: Todas / Enviadas / Fallidas.")
    bullet(doc, "Filtro por tipo de evento.")

    heading(doc, "Tabla de notificaciones", level=2)
    table_simple(doc,
        ["Columna", "Descripción"],
        [
            ["Fecha", "Fecha y hora del envío"],
            ["Destinatario", "Nombre y correo del destinatario"],
            ["Evento", "Tipo de evento que disparó la notificación"],
            ["Asunto", "Asunto del correo enviado"],
            ["Estado", "Enviado / Fallido / Pendiente"],
            ["Acciones", "Botón para reenviar notificaciones fallidas"],
        ]
    )
    doc.add_page_break()

    # ── 18 ──
    heading(doc, "18. Exportar datos (CSV)")
    para(doc, "Para exportar la información de proyectos en formato CSV:")
    numbered(doc, 'Vaya a "Proyectos" en el menú lateral.')
    numbered(doc, 'Haga clic en el botón "Exportar CSV".')
    numbered(doc, "Se descargará un archivo con la información de todos los proyectos incluyendo: título, estudiantes, estado, evaluadores, fechas y calificaciones.")
    doc.add_page_break()

    # ── 19 ──
    heading(doc, "19. Preguntas frecuentes")

    heading(doc, "¿Puedo gestionar proyectos de cualquier programa?", level=2)
    para(doc, "No. Solo puede gestionar proyectos de los programas que le fueron asignados por el superadministrador.")

    heading(doc, "¿Cómo reenvío una notificación fallida?", level=2)
    para(doc, 'Vaya a Notificaciones, filtre por "Fallidas" y haga clic en "Reenviar" junto a la notificación que desea reintentar.')

    heading(doc, "¿Puedo cambiar la nota final manualmente?", level=2)
    para(doc, "Sí. En el detalle del proyecto, sección de calificaciones, existe un campo para sobrescribir la nota final ponderada. Use esta función con precaución.")

    heading(doc, "¿Qué pasa si un evaluador no responde?", level=2)
    para(doc, "Puede reemplazar al evaluador desde el detalle del proyecto. El nuevo evaluador recibirá la notificación y podrá evaluar en su lugar.")

    heading(doc, "Flujo completo de un proyecto", level=2)
    para(doc, "A continuación se describe el flujo completo de estados de un proyecto de grado:")
    numbered(doc, "Borrador → Estudiante crea y guarda el proyecto.")
    numbered(doc, "Enviado → Estudiante envía a evaluación.")
    numbered(doc, "Revisión administrativa → Usted revisa la documentación.")
    numbered(doc, "Evaluadores asignados → Usted asigna 2 evaluadores.")
    numbered(doc, "En evaluación académica → Evaluadores trabajan en la revisión.")
    numbered(doc, "Concepto emitido → Evaluadores envían calificaciones.")
    numbered(doc, "Correcciones solicitadas → (Si aplica) Estudiante corrige y reenvía.")
    numbered(doc, "Aprobada para sustentación → Evaluaciones satisfactorias.")
    numbered(doc, "Sustentación programada → Usted agenda la fecha de defensa.")
    numbered(doc, "Sustentación evaluada → Evaluadores califican la presentación.")
    numbered(doc, "Acta firmada → Todas las firmas digitales recolectadas.")
    numbered(doc, "Finalizado → Proyecto de grado completado.")

    save(doc, "Manual_Profesor_Administrador.docx")


# ═══════════════════════════════════════════════════════════════════
#  MANUAL 4: SUPERADMIN
# ═══════════════════════════════════════════════════════════════════

def manual_superadmin():
    doc = new_doc()
    title_page(doc, "Manual de Usuario", "Rol: Superadministrador")

    heading(doc, "Tabla de Contenido")
    toc = [
        "1. Introducción",
        "2. Inicio de sesión",
        "3. Panel de administración",
        "4. Todas las funciones del Profesor Administrador",
        "5. Gestión de usuarios",
        "6. Configuración de SMTP (correo electrónico)",
        "7. Reglas y plantillas de notificación",
        "8. Configuración global de pesos",
        "9. Administración global de programas",
        "10. Preguntas frecuentes",
    ]
    for item in toc:
        bullet(doc, item)
    doc.add_page_break()

    # ── 1 ──
    heading(doc, "1. Introducción")
    para(doc, "SisTesis es el Sistema de Gestión de Proyectos de Grado de la Facultad de Ingeniería de la Universidad de San Buenaventura – Cali. Este manual describe el uso de la plataforma desde el rol de Superadministrador.")
    para(doc, "El superadministrador tiene acceso completo a todas las funcionalidades del sistema, incluyendo todas las funciones del Profesor Administrador más funciones exclusivas de configuración global.")
    para(doc, "Funciones exclusivas del superadministrador:")
    bullet(doc, "Gestión de todos los usuarios del sistema (estudiantes, evaluadores, administradores).")
    bullet(doc, "Configuración del servidor de correo electrónico (SMTP).")
    bullet(doc, "Personalización de reglas y plantillas de notificación.")
    bullet(doc, "Configuración global de pesos de evaluación.")
    bullet(doc, "Acceso a todos los programas y proyectos sin restricción.")
    doc.add_page_break()

    # ── 2 ──
    heading(doc, "2. Inicio de sesión")
    numbered(doc, f"Ingrese a {URL}/login/staff en su navegador.")
    numbered(doc, "Escriba su correo de superadministrador y contraseña.")
    numbered(doc, 'Haga clic en "Iniciar sesión".')
    numbered(doc, "Será redirigido al panel de administración con acceso completo.")
    para(doc, "")
    warning_box(doc, "Las credenciales del superadministrador deben mantenerse seguras. Cambie la contraseña por defecto en el primer ingreso.")
    doc.add_page_break()

    # ── 3 ──
    heading(doc, "3. Panel de administración")
    para(doc, "El panel de administración del superadministrador es idéntico al del Profesor Administrador, con la diferencia de que muestra estadísticas de TODOS los programas del sistema, no solo los asignados.")
    para(doc, "Además del menú estándar, el superadministrador tiene acceso a las siguientes opciones adicionales en el menú lateral:")

    table_simple(doc,
        ["Opción adicional", "Descripción"],
        [
            ["Usuarios", "Gestión completa de cuentas de usuario"],
            ["Configuración SMTP", "Configuración del servidor de correo"],
            ["Reglas de notificación", "Personalización de plantillas de email"],
        ]
    )
    para(doc, "")
    note_box(doc, "Consulte el Manual del Profesor Administrador para las funciones compartidas: gestión de proyectos, evaluadores, programas, rúbricas, elementos de revisión, pesos, notificaciones y exportación CSV.")
    doc.add_page_break()

    # ── 4 ──
    heading(doc, "4. Todas las funciones del Profesor Administrador")
    para(doc, "El superadministrador tiene acceso a todas las funciones descritas en el Manual del Profesor Administrador, incluyendo:")
    bullet(doc, "Gestión y revisión de proyectos de grado (de TODOS los programas).")
    bullet(doc, "Asignación, reemplazo y eliminación de evaluadores.")
    bullet(doc, "Seguimiento de evaluaciones y vencimientos.")
    bullet(doc, "Gestión de correcciones y rondas de revisión.")
    bullet(doc, "Programación de sustentaciones.")
    bullet(doc, "Gestión de actas y firmas digitales.")
    bullet(doc, "Finalización de proyectos.")
    bullet(doc, "Gestión de evaluadores y programas.")
    bullet(doc, "Configuración de rúbricas y checklist.")
    bullet(doc, "Configuración de pesos de evaluación por programa.")
    bullet(doc, "Historial de notificaciones y reenvío.")
    bullet(doc, "Exportación de datos en CSV.")
    para(doc, "")
    para(doc, "La diferencia clave es que el superadministrador no tiene restricción por programa: puede ver y gestionar TODOS los proyectos, evaluadores y configuraciones del sistema.", bold=True)
    doc.add_page_break()

    # ── 5 ──
    heading(doc, "5. Gestión de usuarios")
    para(doc, 'Desde el menú lateral, seleccione "Usuarios" para acceder a la gestión completa de cuentas.')

    heading(doc, "Vista principal", level=2)
    para(doc, "Tabla con todos los usuarios del sistema:")
    table_simple(doc,
        ["Columna", "Descripción"],
        [
            ["Nombre", "Nombre completo del usuario"],
            ["Correo", "Correo electrónico institucional"],
            ["Cédula", "Número de documento de identidad"],
            ["Roles", "Roles asignados (estudiante, evaluador, admin, superadmin)"],
            ["Programas", "Programas asignados (para administradores)"],
            ["Acciones", "Editar, eliminar, enviar credenciales"],
        ]
    )

    heading(doc, "Crear nuevo usuario", level=2)
    numbered(doc, 'Haga clic en "Crear usuario".')
    numbered(doc, "Complete el formulario:")
    table_simple(doc,
        ["Campo", "Obligatorio", "Descripción"],
        [
            ["Nombre completo", "Sí", "Nombre y apellidos del usuario"],
            ["Código de estudiante", "Solo estudiantes", "Código institucional"],
            ["Cédula", "Sí", "Número de documento de identidad"],
            ["Correo institucional", "Sí", "Correo electrónico (@usbcali.edu.co)"],
            ["Rol(es)", "Sí", "Seleccione: estudiante, evaluador, admin, superadmin"],
            ["Programas", "Solo admin", "Programas que administrará"],
            ["Contraseña", "Sí", "Contraseña inicial del usuario"],
        ]
    )
    numbered(doc, 'Haga clic en "Guardar".')
    numbered(doc, "El sistema creará la cuenta y opcionalmente enviará credenciales por correo.")

    heading(doc, "Editar usuario", level=2)
    para(doc, "Haga clic en el botón de edición junto al usuario. Puede modificar todos los campos incluyendo roles y programas asignados.")

    heading(doc, "Eliminar usuario", level=2)
    para(doc, "Haga clic en el botón de eliminar. Se le pedirá confirmación antes de proceder.")
    warning_box(doc, "Eliminar un usuario es permanente. Verifique que el usuario no tenga proyectos o evaluaciones activas antes de eliminarlo.")

    heading(doc, "Enviar credenciales", level=2)
    para(doc, 'Haga clic en "Enviar credenciales" para reenviar el correo con usuario y contraseña al usuario seleccionado.')
    doc.add_page_break()

    # ── 6 ──
    heading(doc, "6. Configuración de SMTP (correo electrónico)")
    para(doc, 'Desde el menú lateral, seleccione "Configuración SMTP".')
    para(doc, "Esta sección permite configurar el servidor de correo electrónico que el sistema usará para enviar todas las notificaciones.")

    heading(doc, "Campos de configuración", level=2)
    table_simple(doc,
        ["Campo", "Ejemplo", "Descripción"],
        [
            ["Host SMTP", "smtp.gmail.com", "Dirección del servidor de correo"],
            ["Puerto", "587", "Puerto del servidor (587 para TLS, 465 para SSL)"],
            ["Usuario", "sistesis@usbcali.edu.co", "Cuenta de correo para autenticación"],
            ["Contraseña", "********", "Contraseña o contraseña de aplicación"],
            ["Cifrado", "TLS", "Tipo de cifrado: TLS o SSL"],
            ["Es predeterminado", "Sí", "Marcar como configuración activa del sistema"],
        ]
    )

    heading(doc, "Acciones", level=2)
    bullet(doc, "Probar conexión: Verifica que los datos del servidor son correctos.")
    bullet(doc, "Enviar correo de prueba: Envía un correo de ejemplo para verificar el funcionamiento completo.")
    bullet(doc, "Guardar: Almacena la configuración.")
    para(doc, "")
    warning_box(doc, "Si usa Gmail, necesitará una 'Contraseña de aplicación' en lugar de la contraseña normal de la cuenta. Consulte la documentación de Google para generarla.")
    doc.add_page_break()

    # ── 7 ──
    heading(doc, "7. Reglas y plantillas de notificación")
    para(doc, 'Desde el menú lateral, seleccione "Reglas de notificación".')
    para(doc, "Aquí puede configurar qué correos se envían, a quién y con qué contenido para cada evento del sistema.")

    heading(doc, "Tipos de eventos configurables", level=2)
    table_simple(doc,
        ["Evento", "Descripción"],
        [
            ["Tesis enviada", "Un estudiante envía su proyecto a evaluación"],
            ["Retroalimentación administrativa", "El administrador envía comentarios al estudiante"],
            ["Evaluadores asignados", "Se asignan evaluadores a un proyecto"],
            ["Evaluación enviada", "Un evaluador completa su evaluación"],
            ["Sustentación programada", "Se agenda la fecha de defensa"],
            ["Firma de acta", "Se requiere una firma digital"],
            ["Cambio de estado", "El estado del proyecto cambia"],
            ["Evaluador removido/reemplazado", "Se cambia un evaluador en la asignación"],
        ]
    )

    heading(doc, "Configuración por evento", level=2)
    para(doc, "Para cada evento puede configurar:")

    heading(doc, "A. Destinatarios", level=3)
    para(doc, "Active o desactive las casillas según quién debe recibir la notificación:")
    bullet(doc, "Estudiante(s) del proyecto.")
    bullet(doc, "Administrador(es) del programa.")
    bullet(doc, "Evaluador(es) asignados.")

    heading(doc, "B. Plantilla de correo", level=3)
    para(doc, "Personalice el contenido del correo para cada evento:")
    bullet(doc, "Asunto: Línea de asunto del correo.")
    bullet(doc, "Cuerpo HTML: Contenido del correo con formato enriquecido.")

    heading(doc, "Variables disponibles en plantillas", level=2)
    table_simple(doc,
        ["Variable", "Descripción"],
        [
            ["{{destinatario_nombre}}", "Nombre del destinatario"],
            ["{{titulo_tesis}}", "Título del proyecto de grado"],
            ["{{nombres_estudiantes}}", "Nombre(s) de los estudiantes"],
            ["{{correos_estudiantes}}", "Correo(s) de los estudiantes"],
            ["{{nombres_evaluadores}}", "Nombre(s) de los evaluadores"],
            ["{{programa}}", "Nombre del programa académico"],
            ["{{fecha}}", "Fecha actual"],
            ["{{fecha_sustentacion}}", "Fecha de la sustentación"],
            ["{{lugar_sustentacion}}", "Lugar de la sustentación"],
            ["{{descripcion}}", "Descripción del evento"],
        ]
    )
    para(doc, "Inserte las variables en el cuerpo del correo usando la sintaxis de doble llave. El sistema las reemplazará por los valores reales al enviar cada notificación.")
    doc.add_page_break()

    # ── 8 ──
    heading(doc, "8. Configuración global de pesos")
    para(doc, "El superadministrador puede establecer los pesos de evaluación por defecto para todo el sistema.")
    para(doc, "Los programas individuales pueden sobrescribir estos valores, pero si no lo hacen, usarán los valores por defecto establecidos aquí.")

    table_simple(doc,
        ["Componente", "Peso por defecto", "Descripción"],
        [
            ["Documento", "70%", "Porcentaje de la nota final correspondiente al documento"],
            ["Sustentación", "30%", "Porcentaje de la nota final correspondiente a la sustentación oral"],
        ]
    )
    doc.add_page_break()

    # ── 9 ──
    heading(doc, "9. Administración global de programas")
    para(doc, "A diferencia del Profesor Administrador que solo ve sus programas asignados, el superadministrador puede:")
    bullet(doc, "Ver y editar TODOS los programas del sistema.")
    bullet(doc, "Crear nuevos programas y asignar administradores a cada uno.")
    bullet(doc, "Modificar ventanas de recepción de cualquier programa.")
    bullet(doc, "Asignar o remover profesores administradores de cualquier programa.")
    bullet(doc, "Eliminar programas (con confirmación).")
    bullet(doc, "Configurar rúbricas, elementos de revisión y pesos para cualquier programa.")
    doc.add_page_break()

    # ── 10 ──
    heading(doc, "10. Preguntas frecuentes")

    heading(doc, "¿Cómo creo una cuenta de administrador para un profesor?", level=2)
    numbered(doc, 'Vaya a "Usuarios" en el menú lateral.')
    numbered(doc, 'Haga clic en "Crear usuario".')
    numbered(doc, 'Seleccione el rol "admin".')
    numbered(doc, "Asigne los programas que administrará.")
    numbered(doc, "Guarde y el profesor recibirá sus credenciales por correo.")

    heading(doc, "¿Cómo cambio la contraseña por defecto del superadmin?", level=2)
    para(doc, "Edite su propio usuario desde la sección de Usuarios. Ingrese la nueva contraseña y guarde los cambios.")

    heading(doc, "¿Qué pasa si el servidor SMTP no funciona?", level=2)
    para(doc, 'Las notificaciones quedarán en estado "Fallido". Corrija la configuración SMTP y use el botón "Reenviar" en el historial de notificaciones para reintentar los envíos fallidos.')

    heading(doc, "¿Los cambios en plantillas de notificación aplican retroactivamente?", level=2)
    para(doc, "No. Los cambios en las plantillas solo aplican a notificaciones futuras. Las notificaciones ya enviadas no se ven afectadas.")

    heading(doc, "¿Puedo tener varios superadministradores?", level=2)
    para(doc, "Sí. Puede crear múltiples usuarios con rol de superadministrador desde la sección de Usuarios.")

    heading(doc, "Resumen de acceso por rol", level=2)
    table_simple(doc,
        ["Funcionalidad", "Estudiante", "Evaluador", "Admin", "Superadmin"],
        [
            ["Registrar proyecto", "✓", "✗", "✗", "✗"],
            ["Enviar proyecto", "✓", "✗", "✗", "✗"],
            ["Editar proyecto (borrador)", "✓", "✗", "✗", "✗"],
            ["Enviar correcciones", "✓", "✗", "✗", "✗"],
            ["Evaluar documento", "✗", "✓", "✗", "✗"],
            ["Evaluar sustentación", "✗", "✓", "✗", "✗"],
            ["Firmar acta", "✓ (token)", "✓", "✓", "✓"],
            ["Ver timeline", "✓", "✓", "✓", "✓"],
            ["Asignar evaluadores", "✗", "✗", "✓", "✓"],
            ["Gestionar evaluadores", "✗", "✗", "✓", "✓"],
            ["Configurar rúbricas", "✗", "✗", "✓", "✓"],
            ["Gestionar programas", "✗", "✗", "✓ (propios)", "✓ (todos)"],
            ["Gestionar usuarios", "✗", "✗", "✗", "✓"],
            ["Configurar SMTP", "✗", "✗", "✗", "✓"],
            ["Reglas de notificación", "✗", "✗", "✗", "✓"],
            ["Ver reportes", "✗", "✗", "✓", "✓"],
            ["Exportar CSV", "✗", "✗", "✓", "✓"],
        ]
    )

    save(doc, "Manual_Superadministrador.docx")


# ═══════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generando manuales de SisTesis...\n")
    manual_estudiante()
    manual_evaluador()
    manual_admin()
    manual_superadmin()
    print(f"\n✅ Manuales generados en: {OUTPUT_DIR}")
